import base64
import io

import docx
import pytest
from docx.oxml.ns import qn
from PIL import Image

from app import quotation, quotation_docx
from app.quotation import QuotationCustomer, QuotationItem, QuotationProposal
from app.quotation_companies import get_company

# Minimal 1x1 red PNG, used to verify image embedding doesn't crash the render.
_PNG_B64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="


def _webp_data_url(size=(120, 80), color=(30, 80, 160)):
    buf = io.BytesIO()
    Image.new("RGB", size, color=color).save(buf, format="WEBP")
    return "data:image/webp;base64," + base64.b64encode(buf.getvalue()).decode()


def _count_drawings(cell):
    return sum(len(p._element.findall(".//" + qn("w:drawing"))) for p in cell.paragraphs)


def _customer(**overrides):
    base = dict(customer_name="Samaj Dental Hospital", designation="The Director", company_name="Samaj Dental Hospital")
    base.update(overrides)
    return QuotationCustomer(**base)


def _proposal(**overrides):
    base = dict(title="Proposal for RVG Sensor & Portable X-Ray", subject="Quotation for RVG Sensor", quotation_date="2026-07-22")
    base.update(overrides)
    return QuotationProposal(**base)


def _render_and_reopen(items, customer=None, proposal=None):
    company = get_company("garg_dental")
    totals = quotation.compute_totals(items, vat_rate=company.default_vat_rate)
    content = quotation_docx.render_quotation_docx(company, customer or _customer(), proposal or _proposal(), items, totals)
    return docx.Document(io.BytesIO(content)), totals


def test_renders_one_item_with_customer_and_product_details():
    items = [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1, brand="Woodpecker & Co", model="H1")]
    doc, totals = _render_and_reopen(items, customer=_customer(company_name="Smith & Sons Clinic"))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Smith & Sons Clinic" in full_text
    assert "Proposal for RVG Sensor & Portable X-Ray" in full_text

    table = doc.tables[0]
    description_cell = table.rows[1].cells[1].text
    assert "RVG Sensor" in description_cell
    assert "Woodpecker & Co" in description_cell
    assert "108,000.00" in table.rows[1].cells[2].text


def test_renders_multiple_items_as_separate_rows():
    items = [
        QuotationItem(product_name="RVG Sensor", price=108000, quantity=1),
        QuotationItem(product_name="Portable X-Ray", price=175000, quantity=1),
        QuotationItem(product_name="Autoclave", price=50000, quantity=2, discount_percent=10),
    ]
    doc, totals = _render_and_reopen(items)
    table = doc.tables[0]
    # header row + 3 item rows + 4 totals rows
    assert len(table.rows) == 8
    assert table.rows[1].cells[1].text.startswith("RVG Sensor")
    assert table.rows[2].cells[1].text.startswith("Portable X-Ray")
    assert table.rows[3].cells[1].text.startswith("Autoclave")

    full_text = "\n".join(c.text for row in table.rows for c in row.cells)
    assert f"{totals.grand_total:,.2f}" in full_text


def test_renders_with_no_items_without_crashing():
    doc, totals = _render_and_reopen([])
    table = doc.tables[0]
    # header row + 0 item rows + 4 totals rows
    assert len(table.rows) == 5
    assert totals.grand_total == 0


def test_renders_product_image_without_crashing():
    data_url = "data:image/png;base64," + _PNG_B64
    items = [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1, image=data_url)]
    doc, _ = _render_and_reopen(items)
    table = doc.tables[0]
    # header row + 1 item row + 4 totals rows
    assert len(table.rows) == 6


def test_renders_features_and_optional_fields():
    items = [
        QuotationItem(
            product_name="RVG Sensor",
            price=108000,
            quantity=1,
            features=["Ultra-high resolution CMOS sensor.", "Slim design."],
            warranty="24 months",
            specifications=["25 lp/mm"],
            accessories=["USB cable"],
            installation_notes="On-site setup.",
        )
    ]
    doc, _ = _render_and_reopen(items)
    description_cell = doc.tables[0].rows[1].cells[1].text
    assert "Ultra-high resolution CMOS sensor." in description_cell
    assert "Warranty: 24 months" in description_cell
    assert "25 lp/mm" in description_cell
    assert "USB cable" in description_cell
    assert "On-site setup." in description_cell


def test_optional_customer_notes_appear_when_present():
    doc, _ = _render_and_reopen(
        [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1)],
        customer=_customer(notes="Please expedite delivery."),
    )
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Please expedite delivery." in full_text


def test_unknown_company_raises():
    from app.exceptions import UnknownCompanyError

    with pytest.raises(UnknownCompanyError):
        get_company("acme_dental")


def test_mrp_line_present_when_set():
    items = [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1, mrp=130000)]
    doc, _ = _render_and_reopen(items)
    description_cell = doc.tables[0].rows[1].cells[1].text
    assert "MRP: 130,000.00" in description_cell


def test_mrp_line_absent_when_unset():
    items = [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1)]
    doc, _ = _render_and_reopen(items)
    description_cell = doc.tables[0].rows[1].cells[1].text
    assert "MRP:" not in description_cell


def test_webp_image_is_converted_and_embedded():
    items = [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1, image=_webp_data_url())]
    doc, _ = _render_and_reopen(items)
    cell = doc.tables[0].rows[1].cells[1]
    assert _count_drawings(cell) == 1

    # Every embedded media part must be a real PNG (magic bytes), regardless
    # of the source format the browser sent - this is what makes the image
    # reliably open in any Word version.
    for part in doc.part.package.iter_parts():
        partname = str(part.partname)
        if "media/image" in partname and not partname.lower().endswith(".jpeg"):
            assert part.blob[:8] == b"\x89PNG\r\n\x1a\n", f"{partname} is not a real PNG"


def test_multi_item_render_with_mixed_image_formats_and_mrp():
    items = [
        QuotationItem(product_name="RVG Sensor", price=108000, quantity=1, mrp=130000, image=_webp_data_url()),
        QuotationItem(product_name="Portable X-Ray", price=175000, quantity=1),
        QuotationItem(product_name="Autoclave", price=50000, quantity=2, discount_percent=10, mrp=60000),
    ]
    doc, totals = _render_and_reopen(items)
    table = doc.tables[0]
    assert len(table.rows) == 8
    assert "MRP: 130,000.00" in table.rows[1].cells[1].text
    assert "MRP:" not in table.rows[2].cells[1].text
    assert "MRP: 60,000.00" in table.rows[3].cells[1].text
    assert _count_drawings(table.rows[1].cells[1]) == 1
    assert _count_drawings(table.rows[2].cells[1]) == 0


def test_corrupt_image_omitted_without_crashing():
    items = [QuotationItem(product_name="RVG Sensor", price=108000, quantity=1, image="data:image/png;base64,bm90IGFuIGltYWdl")]
    doc, _ = _render_and_reopen(items)
    cell = doc.tables[0].rows[1].cells[1]
    assert _count_drawings(cell) == 0
