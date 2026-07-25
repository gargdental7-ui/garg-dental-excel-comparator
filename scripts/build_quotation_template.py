"""One-off build script: turns the reference Garg Dental proposal .docx into
the docxtpl template used by app/quotation_docx.py. Run again only if the
reference document or the desired tag layout changes:

    source .venv/bin/activate
    python3 scripts/build_quotation_template.py <path-to-reference.docx>

Every edit below replaces run/cell TEXT with Jinja tags in place, reusing
the original run's XML (font, bold, color, size) so formatting is
untouched - this is the "docxtpl on the real reference file" approach: the
template is the reference file, minimally edited, not a reconstruction.
"""
import copy
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Mm

OUTPUT_PATH = Path(__file__).parent.parent / "app" / "quotation_templates" / "equipment_proposal_garg_dental.docx"


def set_run_text(run, text):
    run.text = text


def clear_other_runs(paragraph, keep_index=0, skip_indexes=()):
    keep_indexes = {keep_index, *skip_indexes}
    for i, run in enumerate(paragraph.runs):
        if i not in keep_indexes:
            set_run_text(run, "")


def strip_drawings(run):
    for drawing in run._r.findall(qn("w:drawing")):
        run._r.remove(drawing)


def strip_numpr(paragraph):
    """Remove bullet-list numbering from `paragraph`, if any. Used on
    {%p for%}/{%p endfor%}/{%p if%}/{%p endif%} tag paragraphs and on
    label lines that were cloned from a bulleted line - those paragraphs
    still render an (empty) bullet glyph even though their text collapses
    to nothing, unless the numPr is explicitly removed."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def apply_numpr(paragraph, numpr_template):
    """Attach a copy of a previously-captured <w:numPr> to `paragraph`,
    making it a bulleted list item with the same numbering definition as
    the reference document's feature bullets."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(copy.deepcopy(numpr_template))


def set_cant_split(row):
    """Set w:cantSplit on a table row so it moves to the next page as a
    whole instead of being cut mid-content at the page break."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))


def clone_paragraph_after(paragraph):
    """Insert a copy of `paragraph`'s XML (same pPr/rPr styling) right
    after it, returning the new Paragraph wrapper. Used for optional
    product-detail lines that don't exist in the reference sample."""
    new_p_element = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p_element)
    return docx.text.paragraph.Paragraph(new_p_element, paragraph._parent)


def set_paragraph_single_run_text(paragraph, text):
    """Collapse a paragraph down to a single run holding `text`, reusing
    the paragraph's first run's formatting (font/bold/etc)."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    set_run_text(paragraph.runs[0], text)
    clear_other_runs(paragraph, keep_index=0)


def edit_letterhead_and_title(doc):
    title_p = doc.paragraphs[6]
    set_paragraph_single_run_text(title_p, "{{ proposal.title }}")


def edit_submitted_to_block(doc):
    # Paragraph 14: "The Director\nSamaj Dental Hospital" - run[1] is a
    # real <w:br/> line-break run (verified against the source XML), not
    # literal text, so it's left untouched; only the two text runs around
    # it change.
    p14 = doc.paragraphs[14]
    set_run_text(p14.runs[0], "{{ customer.designation }}")
    set_run_text(p14.runs[2], "{{ customer.company_name }}")
    # keep run[1] (the real <w:br/> line-break run) completely untouched
    clear_other_runs(p14, keep_index=0, skip_indexes=(1, 2))

    p15 = doc.paragraphs[15]  # "Bibhuti Janak Marg, Kathmandu"
    set_paragraph_single_run_text(p15, "{{ customer.address }}")
    # paragraph 16 ("Nepal") stays static - Garg Dental serves Nepal only.


def edit_ref_and_date(doc):
    p33 = doc.paragraphs[33]  # "Ref No: GD/Q/SDH/2082/83-01"
    set_run_text(p33.runs[0], "Ref No: {{ customer.reference_number }}")
    clear_other_runs(p33, keep_index=0)

    p34 = doc.paragraphs[34]  # "Date: 28th June, 2026"
    set_run_text(p34.runs[0], "Date: {{ proposal.quotation_date }}")
    clear_other_runs(p34, keep_index=0)


def edit_to_address_block(doc):
    p37 = doc.paragraphs[37]  # "The Director"
    set_paragraph_single_run_text(p37, "{{ customer.designation }}")

    p38 = doc.paragraphs[38]  # "Samaj Dental Hospital"
    set_paragraph_single_run_text(p38, "{{ customer.company_name }}")

    p39 = doc.paragraphs[39]  # "Bibhuti Janak Marg, Kathmandu"
    set_paragraph_single_run_text(p39, "{{ customer.address }}")
    # paragraph 40 ("Nepal") stays static.


def edit_subject_and_intro(doc):
    p41 = doc.paragraphs[41]  # "Subject: Quotation for ..."
    set_run_text(p41.runs[0], "Subject: {{ proposal.subject }}")
    clear_other_runs(p41, keep_index=0)

    # Intro paragraph 43: keep Garg Dental's boilerplate wording, append an
    # optional customer-notes line to its last run rather than inserting a
    # whole new paragraph (lower risk of corrupting the XML tree).
    p43 = doc.paragraphs[43]
    last_run = p43.runs[-1]
    set_run_text(
        last_run,
        last_run.text + "{{ (' Notes: ' + customer.notes) if customer.notes else '' }}",
    )


def copy_pPr(paragraph):
    pPr = paragraph._p.find(qn("w:pPr"))
    return copy.deepcopy(pPr) if pPr is not None else None


def copy_rPr(run):
    rPr = run._r.find(qn("w:rPr"))
    return copy.deepcopy(rPr) if rPr is not None else None


def apply_pPr(paragraph, pPr_template):
    p = paragraph._p
    existing = p.find(qn("w:pPr"))
    if existing is not None:
        p.remove(existing)
    if pPr_template is not None:
        p.insert(0, copy.deepcopy(pPr_template))


def apply_rPr(run, rPr_template):
    r = run._r
    existing = r.find(qn("w:rPr"))
    if existing is not None:
        r.remove(existing)
    if rPr_template is not None:
        r.insert(0, copy.deepcopy(rPr_template))


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "nil"}))
    tblPr.append(borders)


def edit_terms_and_conditions(doc):
    # Paragraphs 48-56 are the 9 static "Label:\t\tText" lines, tab-aligned.
    # A single tab only lines up the FIRST line of a wrapped value -
    # continuation lines fall back to the paragraph's left margin (visible
    # in the rendered output, e.g. "Shipment From:"). Replace the whole
    # block with a real borderless 2-column table so every wrapped line
    # stays indented under the value column - the sub-bullet
    # warranty-invalidity list (57-63) stays untouched/static.
    p48 = doc.paragraphs[48]
    pPr_template = copy_pPr(p48)
    rPr_template = copy_rPr(p48.runs[0])
    for p in doc.paragraphs[49:57]:
        p._p.getparent().remove(p._p)

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Mm(38)
        row.cells[1].width = Mm(122)

    def style_cell(cell, text):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        apply_pPr(paragraph, pPr_template)
        apply_rPr(run, rPr_template)

    # docxtpl's {%tr for %}/{%tr endfor %} tags each collapse their own
    # ENTIRE row down to the bare tag (same mechanism as the product
    # table), so the loop needs a dedicated for-row, this content row, and
    # a dedicated endfor-row.
    for_row, content_row, endfor_row = table.rows
    style_cell(for_row.cells[0], "{%tr for term in company.terms %}")
    style_cell(content_row.cells[0], "{{ term.0 }}:")
    style_cell(content_row.cells[1], "{{ term.1 }}")
    style_cell(endfor_row.cells[0], "{%tr endfor %}")
    set_cant_split(content_row)
    remove_table_borders(table)

    # doc.add_table() appends at the end of the body - relocate the table
    # to where the old tab-aligned paragraphs used to live, then drop p48.
    tbl_element = table._tbl
    tbl_element.getparent().remove(tbl_element)
    p48._p.addprevious(tbl_element)
    p48._p.getparent().remove(p48._p)


def edit_product_table(doc):
    table = doc.tables[0]
    row = table.rows[1]
    cells = row.cells

    # docxtpl's {%tr for %}/{%tr endfor %} tags each collapse their own
    # ENTIRE <w:tr> down to the bare tag (verified empirically - the tag
    # can't share a row with the content it's supposed to repeat, or the
    # content gets eaten along with it). So the loop needs three rows: a
    # dedicated "for" marker row, this untouched content row, and a
    # dedicated "endfor" marker row - both markers vanish on render,
    # leaving only the repeated content row(s).
    sn_p = cells[0].paragraphs[0]
    set_paragraph_single_run_text(sn_p, "{{ loop.index }}.")

    # --- DESCRIPTION column: rewrite each paragraph in place ---
    desc_paragraphs = cells[1].paragraphs
    set_paragraph_single_run_text(desc_paragraphs[0], "{{ item.product_name }}")
    set_paragraph_single_run_text(desc_paragraphs[1], "MODEL : {{ item.model }}")
    set_paragraph_single_run_text(desc_paragraphs[2], "BRAND: {{ item.brand }}")
    set_paragraph_single_run_text(desc_paragraphs[3], "ORIGIN: {{ item.origin }}")

    # paragraph 4 held the reference's static product image - strip the
    # drawing and replace with the per-item conditional image tag.
    image_p = desc_paragraphs[4]
    image_run = image_p.runs[0] if image_p.runs else image_p.add_run("")
    strip_drawings(image_run)
    set_run_text(image_run, "{% if item.image %}{{ item.image }}{% endif %}")
    clear_other_runs(image_p, keep_index=0)

    # Capture the pristine bullet numbering (numId) from the reference's
    # first feature line before anything below touches it, so it can be
    # reapplied explicitly to only the paragraphs that are real list items.
    numpr_template = copy.deepcopy(desc_paragraphs[6]._p.find(qn("w:pPr")).find(qn("w:numPr")))

    # "Key Features :" is only meaningful when there ARE features - make it
    # a {%p if %} block (like the optional lines below) instead of a static
    # heading, so an item with no features doesn't show a heading over
    # nothing.
    heading_p = desc_paragraphs[5]
    set_paragraph_single_run_text(heading_p, "{%p if item.features %}")
    strip_numpr(heading_p)
    anchor = clone_paragraph_after(heading_p)
    set_paragraph_single_run_text(anchor, "Key Features :")
    strip_numpr(anchor)
    anchor = clone_paragraph_after(anchor)
    set_paragraph_single_run_text(anchor, "{%p endif %}")
    strip_numpr(anchor)

    # paragraph 6 is the first feature line - reuse its styling for the
    # {%p for %} tag paragraph. {%p for %}/{%p endfor %} each collapse
    # their OWN paragraph down to the bare tag (verified empirically), so
    # for/content/endfor must be three separate paragraphs, not one
    # combined line. Delete the remaining sample feature lines (7, 8, 9)
    # first so paragraph indexes below stay valid.
    for p in desc_paragraphs[7:10]:
        p._p.getparent().remove(p._p)
    set_paragraph_single_run_text(desc_paragraphs[6], "{%p for f in item.features %}")
    strip_numpr(desc_paragraphs[6])
    anchor = clone_paragraph_after(desc_paragraphs[6])
    set_paragraph_single_run_text(anchor, "{{ f }}")
    apply_numpr(anchor, numpr_template)
    anchor = clone_paragraph_after(anchor)
    set_paragraph_single_run_text(anchor, "{%p endfor %}")
    strip_numpr(anchor)

    def append_line(text, bulleted=False):
        nonlocal anchor
        anchor = clone_paragraph_after(anchor)
        set_paragraph_single_run_text(anchor, text)
        if bulleted:
            apply_numpr(anchor, numpr_template)
        else:
            strip_numpr(anchor)

    def append_conditional_line(condition_expr, content_expr):
        # {%p if %}/{%p endif %} collapse their own paragraph AND the
        # paragraph(s) between them when the condition is false (verified
        # empirically) - so a blank/absent field disappears entirely
        # instead of leaving an empty line or bullet behind.
        append_line(f"{{%p if {condition_expr} %}}")
        append_line(content_expr)
        append_line("{%p endif %}")

    def append_loop(list_expr, item_name):
        append_line(f"{{%p for {item_name} in {list_expr} %}}")
        append_line(f"{{{{ {item_name} }}}}", bulleted=True)
        append_line("{%p endfor %}")

    # Optional extra product-detail lines, cloned from the feature-line
    # paragraph's styling (minus the bullet - only actual list items keep
    # numPr), appended after the features loop.
    append_conditional_line("item.mrp_formatted", "MRP: {{ item.mrp_formatted }}")
    append_conditional_line("item.warranty", "Warranty: {{ item.warranty }}")
    append_conditional_line("item.specifications", "Technical Specifications :")
    append_loop("item.specifications", "s")
    append_conditional_line("item.accessories", "Accessories :")
    append_loop("item.accessories", "a")
    append_conditional_line("item.installation_notes", "Installation Notes: {{ item.installation_notes }}")
    append_conditional_line("item.additional_notes", "Additional Notes: {{ item.additional_notes }}")

    # --- RATE column ---
    rate_p = cells[2].paragraphs[0]
    set_paragraph_single_run_text(rate_p, "{{ item.rate_formatted }}")

    # Delete the reference's second sample product row entirely - the
    # content row above now generates one row per real item.
    table.rows[2]._tr.getparent().remove(table.rows[2]._tr)

    # A product's row can run to several lines (image + features + specs);
    # without cantSplit, Word/LibreOffice happily break it mid-paragraph at
    # a page boundary. Force the whole row to move to the next page instead.
    set_cant_split(row)

    # Wrap the content row with dedicated for/endfor marker rows (cloned
    # from it, so column widths/grid stay consistent, then cleared).
    content_tr = row._tr

    for_tr = copy.deepcopy(content_tr)
    content_tr.addprevious(for_tr)

    endfor_tr = copy.deepcopy(content_tr)
    content_tr.addnext(endfor_tr)

    table = doc.tables[0]  # re-fetch: row list changed after XML inserts
    for_row = table.rows[1]
    endfor_row = table.rows[3]
    for_row.cells[0].text = "{%tr for item in items %}"
    for c in for_row.cells[1:]:
        c.text = ""
    endfor_row.cells[0].text = "{%tr endfor %}"
    for c in endfor_row.cells[1:]:
        c.text = ""

    return table


def copy_cell_borders(source_cell, target_cell):
    """Copy a cell's <w:tcBorders> (and vertical alignment) from
    `source_cell` onto `target_cell`. table.add_row() creates brand-new
    cells with no border formatting at all - the table itself has no
    default border style, so without this the row renders with no visible
    grid lines at all, unlike every real product row above it."""
    source_tcPr = source_cell._tc.find(qn("w:tcPr"))
    if source_tcPr is None:
        return
    target_tcPr = target_cell._tc.get_or_add_tcPr()
    for tag in ("w:tcBorders", "w:vAlign"):
        existing = target_tcPr.find(qn(tag))
        if existing is not None:
            target_tcPr.remove(existing)
        source_el = source_tcPr.find(qn(tag))
        if source_el is not None:
            target_tcPr.append(copy.deepcopy(source_el))


def append_totals_rows(table):
    # The real product content row - still carries the reference document's
    # actual per-cell border formatting (see copy_cell_borders).
    style_row = table.rows[2]

    def add_row(label, value_expr, bold=False):
        row = table.add_row()
        row.cells[1].text = label
        row.cells[2].text = value_expr
        for i, cell in enumerate(row.cells):
            copy_cell_borders(style_row.cells[i], cell)
        if bold:
            for cell in (row.cells[1], row.cells[2]):
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    add_row("Subtotal", "{{ totals.subtotal_formatted }}")
    add_row("Discount", "{{ totals.discount_formatted }}")
    add_row("VAT", "{{ totals.vat_formatted }}")
    add_row("Grand Total", "{{ totals.grand_total_formatted }}", bold=True)


def build(reference_path: str):
    doc = docx.Document(reference_path)
    edit_letterhead_and_title(doc)
    edit_submitted_to_block(doc)
    edit_ref_and_date(doc)
    edit_to_address_block(doc)
    edit_subject_and_intro(doc)
    edit_terms_and_conditions(doc)
    table = edit_product_table(doc)
    append_totals_rows(table)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python3 scripts/build_quotation_template.py <reference.docx>")
        sys.exit(1)
    build(sys.argv[1])
